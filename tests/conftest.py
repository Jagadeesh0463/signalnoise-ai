"""
tests/conftest.py

Shared pytest fixtures for all SignalNoise AI tests.
Every test file can use these without importing them — pytest loads this automatically.

Fixtures:
    tmp_txt_file      — creates a temporary .txt file with given content
    tmp_docx_file     — creates a temporary .docx file with given content
    sample_english_text — generates valid English text of given word count
    sample_document   — a ready-to-use Document object
    sample_anon_doc   — a ready-to-use AnonymizedDocument object
    sample_signal     — a ready-to-use Signal object (STRONG)
    sample_weak_signal — a ready-to-use Signal object (WEAK)
    memory_store      — a fresh MemoryStore backed by a temp SQLite DB
    mock_groq         — patches Groq so tests never make real API calls
"""

import os
import uuid
from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

# Set required env vars before any src import
os.environ.setdefault("GROQ_API_KEY", "test-key-conftest")

from src.models import AnonymizedDocument, Document, Evidence, Feedback, Risk, Signal


# ── Text helpers ──────────────────────────────────────────────────────────────

@pytest.fixture
def sample_english_text():
    """
    Factory fixture — call it with a word count to get valid English text.

    Usage:
        def test_something(sample_english_text):
            text = sample_english_text(200)
    """
    def _make(word_count: int = 100) -> str:
        base = (
            "the team reported that the delivery milestone was delayed and the "
            "programme manager escalated to the director of engineering for review "
            "of the outstanding issues with the release pipeline and deployment "
            "the backend lead confirmed the blocker is now on the critical path "
        )
        tokens = base.split()
        result = (tokens * ((word_count // len(tokens)) + 1))[:word_count]
        return " ".join(result)

    return _make


# ── File fixtures ─────────────────────────────────────────────────────────────

@pytest.fixture
def tmp_txt_file(tmp_path, sample_english_text):
    """
    Factory fixture — creates a real .txt file in a temp directory.

    Usage:
        def test_something(tmp_txt_file):
            path = tmp_txt_file("My document content here with enough words " * 5)
    """
    def _make(content: str | None = None, word_count: int = 100) -> Path:
        text = content if content is not None else sample_english_text(word_count)
        file_path = tmp_path / f"test_{uuid.uuid4().hex[:8]}.txt"
        file_path.write_text(text, encoding="utf-8")
        return file_path

    return _make


@pytest.fixture
def tmp_docx_file(tmp_path, sample_english_text):
    """
    Factory fixture — creates a real .docx file in a temp directory.

    Usage:
        def test_something(tmp_docx_file):
            path = tmp_docx_file("My document content here " * 10)
    """
    def _make(content: str | None = None, word_count: int = 100) -> Path:
        from docx import Document as DocxDocument

        text = content if content is not None else sample_english_text(word_count)
        doc = DocxDocument()
        # Split into paragraphs for realistic .docx structure
        for para in text.split("."):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        file_path = tmp_path / f"test_{uuid.uuid4().hex[:8]}.docx"
        doc.save(str(file_path))
        return file_path

    return _make


# ── Model fixtures ────────────────────────────────────────────────────────────

@pytest.fixture
def sample_document(sample_english_text) -> Document:
    """A ready-to-use Document object with realistic content."""
    return Document(
        id=str(uuid.uuid4()),
        filename="sprint_review.txt",
        source_type="meeting_note",
        raw_text=sample_english_text(120),
        word_count=120,
        uploaded_at=datetime.utcnow(),
        processed=False,
    )


@pytest.fixture
def sample_anon_doc(sample_document) -> AnonymizedDocument:
    """A ready-to-use AnonymizedDocument (post-privacy-shield)."""
    return AnonymizedDocument(
        id=str(uuid.uuid4()),
        document_id=sample_document.id,
        anonymized_text=(
            "[Backend-Lead-A] reported that the deployment pipeline has been delayed "
            "for three consecutive weeks. [Platform-Lead-B] confirmed the dependency "
            "on vendor approval is now blocking the release milestone. "
            "[Programme-Manager-C] escalated to the director for urgent review."
        ),
        role_map={
            "John Smith": "[Backend-Lead-A]",
            "Jane Doe": "[Platform-Lead-B]",
            "Priya R": "[Programme-Manager-C]",
        },
        processed_at=datetime.utcnow(),
    )


@pytest.fixture
def sample_signal() -> Signal:
    """A ready-to-use STRONG Signal object."""
    return Signal(
        id=str(uuid.uuid4()),
        title="Repeated deployment pipeline delays",
        category="delivery_risk",
        severity="STRONG",
        confidence_band="high",
        trend="emerging",
        evidence=[
            "deployment pipeline has been delayed for three consecutive weeks",
            "vendor approval is now blocking the release milestone",
        ],
        source_document_ids=[str(uuid.uuid4()), str(uuid.uuid4())],
        suggested_owner_role="Programme-Manager",
        related_teams=["Backend-Team", "Platform-Team"],
        related_projects=[],
        status="active",
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )


@pytest.fixture
def sample_weak_signal() -> Signal:
    """A ready-to-use WEAK Signal object."""
    return Signal(
        id=str(uuid.uuid4()),
        title="Early signs of team capacity strain",
        category="team_health",
        severity="WEAK",
        confidence_band="low",
        trend="emerging",
        evidence=["team reported feeling overwhelmed this sprint"],
        source_document_ids=[str(uuid.uuid4())],
        suggested_owner_role="Delivery-Director",
        related_teams=["Backend-Team"],
        related_projects=[],
        status="active",
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )


@pytest.fixture
def sample_evidence(sample_signal) -> Evidence:
    """A ready-to-use Evidence object linked to sample_signal."""
    return Evidence(
        id=str(uuid.uuid4()),
        signal_id=sample_signal.id,
        document_id=str(uuid.uuid4()),
        snippet="deployment pipeline has been delayed for three consecutive weeks",
        relevance_score=0.91,
        source_count=2,
    )


@pytest.fixture
def sample_risk(sample_signal) -> Risk:
    """A ready-to-use Risk object linked to sample_signal."""
    return Risk(
        id=str(uuid.uuid4()),
        signal_id=sample_signal.id,
        business_impact="Programme delivery milestone at critical risk.",
        confidence_band="high",
        priority="critical",
        suggested_owner_role="Programme-Manager",
        suggested_action="Convene emergency programme review today.",
        root_cause_hypothesis="Resource constraint and unresolved dependency.",
        supporting_document_ids=sample_signal.source_document_ids,
        narration="",
        created_at=datetime.utcnow(),
    )


@pytest.fixture
def sample_feedback(sample_signal) -> Feedback:
    """A ready-to-use Feedback object."""
    return Feedback(
        id=str(uuid.uuid4()),
        signal_id=sample_signal.id,
        reviewer_role="Programme-Manager",
        decision="confirmed",
        comment="Matches what I heard in standup.",
        created_at=datetime.utcnow(),
    )


# ── Store fixture ─────────────────────────────────────────────────────────────

@pytest.fixture
def memory_store(tmp_path):
    """
    A fresh MemoryStore backed by a temporary SQLite database.
    Automatically cleaned up after each test.

    Usage:
        def test_something(memory_store):
            memory_store.save_signal(signal)
    """
    from src.memory.store import MemoryStore
    db_path = tmp_path / "test_signalnoise.db"
    return MemoryStore(db_path=db_path)


# ── Mock fixtures ─────────────────────────────────────────────────────────────

@pytest.fixture
def mock_groq():
    """
    Patches the Groq client so tests never make real API calls.
    Returns a mock that produces a predictable 2-sentence narration.

    Usage:
        def test_narration(mock_groq):
            risk = narrate_risk(risk)
            assert risk.narration  # uses mock response
    """
    mock_response = MagicMock()
    mock_response.choices[0].message.content = (
        "The programme delivery milestone is at critical risk due to a "
        "persistent pipeline blockage. Immediate escalation to the programme "
        "manager is recommended."
    )

    with patch("src.narration.narrator._get_client") as mock_client:
        mock_instance = MagicMock()
        mock_instance.chat.completions.create.return_value = mock_response
        mock_client.return_value = mock_instance
        yield mock_instance
