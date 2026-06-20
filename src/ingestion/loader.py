"""
src/ingestion/loader.py

File Loader — extracts plain text from uploaded documents.
Supports: .txt, .docx, .pdf

Pipeline position:
    Upload → Quality Gate (extension check) → Loader → Quality Gate (text check) → Anonymizer

Usage:
    from src.ingestion.loader import load_file
    from pathlib import Path

    document = load_file(Path("data/raw/meeting_notes.docx"), source_type="meeting_note")
"""

import logging
import uuid
from datetime import datetime
from pathlib import Path

import chardet
import pdfplumber
from docx import Document as DocxDocument

from src.exceptions import LoaderError
from src.ingestion.quality_gate import check, check_file_extension
from src.models import Document

logger = logging.getLogger(__name__)

# Max pages to read from a PDF — prevents zip-bomb / runaway extraction
PDF_MAX_PAGES = 200

VALID_SOURCE_TYPES = {
    "meeting_note",
    "incident_log",
    "ticket",
    "status_report",
}


# ── Public API ────────────────────────────────────────────────────────────────

def load_file(file_path: Path, source_type: str) -> Document:
    """
    Load a file from disk and return a Document ready for the privacy shield.

    Args:
        file_path:   Path to the uploaded file (.txt, .docx, or .pdf)
        source_type: One of: meeting_note | incident_log | ticket | status_report

    Returns:
        Document with raw_text populated and quality gate passed.

    Raises:
        LoaderError: If the file cannot be read, is empty, or fails the quality gate.
    """
    if source_type not in VALID_SOURCE_TYPES:
        raise LoaderError(
            f"Invalid source_type '{source_type}'. "
            f"Must be one of: {', '.join(sorted(VALID_SOURCE_TYPES))}"
        )

    if not file_path.exists():
        raise LoaderError(f"File not found: {file_path}")

    filename = file_path.name

    # Step 1 — extension check before attempting extraction
    ext_result = check_file_extension(filename)
    if ext_result.failed:
        raise LoaderError(f"Rejected at extension check: {ext_result}")

    # Step 2 — extract text based on file type
    ext = file_path.suffix.lower()
    logger.info("Loading file: %s (%s)", filename, ext)

    try:
        if ext == ".txt":
            raw_text = _load_txt(file_path)
        elif ext == ".docx":
            raw_text = _load_docx(file_path)
        elif ext == ".pdf":
            raw_text = _load_pdf(file_path)
        else:
            raise LoaderError(f"Unsupported extension: {ext}")
    except LoaderError:
        raise
    except Exception as exc:
        raise LoaderError(f"Failed to extract text from '{filename}': {exc}") from exc

    # Step 3 — quality gate on extracted text
    quality_result = check(raw_text, filename=filename)
    if quality_result.failed:
        raise LoaderError(
            f"File '{filename}' failed quality gate: {quality_result}"
        )

    if quality_result.warnings:
        for warning in quality_result.warnings:
            logger.warning("Quality warning for '%s': %s", filename, warning)

    # Step 4 — build and return Document
    doc = Document(
        id=str(uuid.uuid4()),
        filename=filename,
        source_type=source_type,
        raw_text=raw_text,
        word_count=quality_result.word_count,
        uploaded_at=datetime.utcnow(),
        processed=False,
    )

    logger.info(
        "Loaded '%s' — %d words, source_type=%s, doc_id=%s",
        filename,
        doc.word_count,
        source_type,
        doc.id[:8],
    )
    return doc


# ── Internal extractors ───────────────────────────────────────────────────────

def _load_txt(file_path: Path) -> str:
    """
    Read a .txt file. Auto-detects encoding using chardet.
    Falls back to UTF-8 with error replacement if detection fails.
    """
    raw_bytes = file_path.read_bytes()

    if not raw_bytes:
        raise LoaderError(f"File is empty: {file_path.name}")

    detected = chardet.detect(raw_bytes)
    encoding = detected.get("encoding") or "utf-8"
    confidence = detected.get("confidence", 0)

    logger.debug(
        "Detected encoding '%s' (confidence %.0f%%) for %s",
        encoding,
        confidence * 100,
        file_path.name,
    )

    try:
        return raw_bytes.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        logger.warning(
            "Encoding '%s' failed for %s — falling back to UTF-8 with replacement.",
            encoding,
            file_path.name,
        )
        return raw_bytes.decode("utf-8", errors="replace")


def _load_docx(file_path: Path) -> str:
    """
    Read a .docx file. Extracts text from paragraphs and tables.
    Skips empty paragraphs.
    """
    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        raise LoaderError(f"Cannot open .docx file '{file_path.name}': {exc}") from exc

    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract text from tables (each cell on its own line)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    if not parts:
        raise LoaderError(f"No text found in .docx file: {file_path.name}")

    return "\n".join(parts)


def _load_pdf(file_path: Path) -> str:
    """
    Read a .pdf file using pdfplumber.
    Limits extraction to PDF_MAX_PAGES to prevent runaway processing.
    Skips image-only pages gracefully.
    """
    parts: list[str] = []

    try:
        with pdfplumber.open(str(file_path)) as pdf:
            total_pages = len(pdf.pages)

            if total_pages > PDF_MAX_PAGES:
                logger.warning(
                    "PDF '%s' has %d pages — reading first %d only.",
                    file_path.name,
                    total_pages,
                    PDF_MAX_PAGES,
                )

            for page in pdf.pages[:PDF_MAX_PAGES]:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text.strip())

    except Exception as exc:
        raise LoaderError(
            f"Cannot read PDF '{file_path.name}': {exc}"
        ) from exc

    if not parts:
        raise LoaderError(
            f"No text extracted from PDF '{file_path.name}'. "
            "The file may be image-only (scanned). Use an OCR tool first."
        )

    return "\n".join(parts)
